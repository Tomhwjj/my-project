#!/usr/bin/env python3
"""
generate_docx.py — 将 Markdown 论文转换为标准 .docx 文件
兼容：Microsoft Office 2016+ / WPS Office 2019+
"""

import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("[ERR] 需要安装 python-docx 库")
    print("      运行: pip install python-docx")
    sys.exit(1)


def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, val in kwargs.items():
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), val.get('val', 'single'))
        element.set(qn('w:sz'), str(val.get('sz', 4)))
        element.set(qn('w:color'), val.get('color', '000000'))
        tcBorders.append(element)
    tcPr.append(tcBorders)


def set_run_font(run, cn_font='宋体', en_font='Times New Roman', size=Pt(12), bold=False):
    """设置run的字体属性"""
    run.font.size = size
    run.bold = bold
    run.font.name = en_font
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rPr.insert(0, rFonts)


def set_paragraph_spacing(paragraph, line_spacing=1.5, space_before=0, space_after=0):
    """设置段落间距"""
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)


def set_first_line_indent(paragraph, chars=2, font_size=12):
    """设置首行缩进（字符数）"""
    paragraph.paragraph_format.first_line_indent = Pt(chars * font_size)


def add_heading_custom(doc, text, level=1, cn_font='黑体', size=None):
    """添加自定义标题"""
    sizes = {0: Pt(18), 1: Pt(16), 2: Pt(14), 3: Pt(12)}
    if size is None:
        size = sizes.get(level, Pt(12))

    p = doc.add_paragraph()
    if level == 0:  # 论文总标题
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, cn_font=cn_font, size=size, bold=True)
    set_paragraph_spacing(p, line_spacing=1.5, space_before=12, space_after=6)
    return p


def add_body_text(doc, text, first_line_indent=True):
    """添加正文段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, cn_font='宋体', size=Pt(12), bold=False)
    set_paragraph_spacing(p, line_spacing=1.5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if first_line_indent:
        set_first_line_indent(p, chars=2)
    return p


def parse_markdown(md_text):
    """简单解析Markdown结构，返回段落列表"""
    lines = md_text.strip().split('\n')
    elements = []
    for line in lines:
        line = line.rstrip()
        if not line:
            elements.append({'type': 'blank'})
        elif line.startswith('# ') or line.startswith('论文标题') or line.startswith('# '):
            elements.append({'type': 'h0', 'text': re.sub(r'^#\s+', '', line)})
        elif line.startswith('## '):
            elements.append({'type': 'h1', 'text': re.sub(r'^##\s+', '', line)})
        elif line.startswith('### '):
            elements.append({'type': 'h2', 'text': re.sub(r'^###\s+', '', line)})
        elif line.startswith('#### '):
            elements.append({'type': 'h3', 'text': re.sub(r'^####\s+', '', line)})
        elif line.startswith('---'):
            elements.append({'type': 'hr'})
        elif re.match(r'^\d+\.\s', line):
            elements.append({'type': 'ordered_list', 'text': re.sub(r'^\d+\.\s+', '', line)})
        elif line.startswith('- ') or line.startswith('* '):
            elements.append({'type': 'bullet', 'text': re.sub(r'^[-*]\s+', '', line)})
        elif line.startswith('> '):
            elements.append({'type': 'quote', 'text': re.sub(r'^>\s+', '', line)})
        elif re.match(r'^\*\*.*\*\*$', line):
            elements.append({'type': 'bold_text', 'text': line.strip('*')})
        else:
            elements.append({'type': 'body', 'text': line})
    return elements


def generate_docx(md_path, output_path, template_path=None):
    """主函数：Markdown → docx"""
    # 读取Markdown
    md_content = Path(md_path).read_text(encoding='utf-8')

    # 创建文档
    if template_path and Path(template_path).exists():
        doc = Document(template_path)
    else:
        doc = Document()

    # 设置默认样式
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    pf = style.paragraph_format
    pf.line_spacing = 1.5

    # 设置页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    # 解析并生成内容
    elements = parse_markdown(md_content)

    skip_indent_types = {'h0', 'h1', 'h2', 'h3', 'hr', 'blank', 'ordered_list', 'bullet', 'quote'}

    for elem in elements:
        t = elem['type']
        text = elem.get('text', '')

        if t == 'h0':
            add_heading_custom(doc, text, level=0)
        elif t == 'h1':
            add_heading_custom(doc, text, level=1)
        elif t == 'h2':
            add_heading_custom(doc, text, level=2)
        elif t == 'h3':
            add_heading_custom(doc, text, level=3)
        elif t == 'body':
            add_body_text(doc, text, first_line_indent=True)
        elif t == 'bold_text':
            p = doc.add_paragraph()
            run = p.add_run(text)
            set_run_font(run, cn_font='宋体', size=Pt(12), bold=True)
            set_paragraph_spacing(p, line_spacing=1.5)
            set_first_line_indent(p, chars=2)
        elif t in ('ordered_list', 'bullet'):
            p = doc.add_paragraph()
            prefix = '• ' if t == 'bullet' else f'{len(doc.paragraphs)}. '
            run = p.add_run(prefix + text)
            set_run_font(run, cn_font='宋体', size=Pt(12))
            set_paragraph_spacing(p, line_spacing=1.5)
        elif t == 'quote':
            p = doc.add_paragraph()
            run = p.add_run(text)
            set_run_font(run, cn_font='楷体', size=Pt(12))
            set_paragraph_spacing(p, line_spacing=1.5)
            p.paragraph_format.left_indent = Cm(1)
        elif t == 'hr':
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 0.5
        # 'blank' is skipped

    # 保存
    doc.save(output_path)
    print(f"[OK] 文档已生成: {output_path}")
    print(f"     格式兼容: Microsoft Office 2016+ / WPS Office 2019+")


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Markdown → 标准论文 .docx 转换器')
    parser.add_argument('--input', '-i', required=True, help='输入的 Markdown 论文文件')
    parser.add_argument('--output', '-o', required=True, help='输出的 .docx 文件路径')
    parser.add_argument('--template', '-t', default=None, help='可选的 Word 模板文件 (.docx)')
    args = parser.parse_args()

    if not Path(args.input).exists():
        print(f"[ERR] 输入文件不存在: {args.input}")
        sys.exit(1)

    generate_docx(args.input, args.output, args.template)
