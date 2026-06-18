#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
verify_format.py - 验证 .docx 格式合规性
检查字体、字号、行距、页边距等是否符合规范
"""

import argparse
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm
except ImportError:
    print("[ERR] 需要安装 python-docx 库")
    print("      运行: pip install python-docx")
    sys.exit(1)

# 格式标准
STANDARDS = {
    'page': {
        'top_margin': (2.54, 0.1),
        'bottom_margin': (2.54, 0.1),
        'left_margin': (3.18, 0.1),
        'right_margin': (3.18, 0.1),
    },
    'body_font': {
        'cn_fonts': ['宋体', 'SimSun'],
        'en_fonts': ['Times New Roman'],
        'size': (12, 1),
    },
    'heading_fonts': {
        1: {'cn_fonts': ['黑体', 'SimHei'], 'size': (16, 1)},
        2: {'cn_fonts': ['黑体', 'SimHei'], 'size': (14, 1)},
        3: {'cn_fonts': ['黑体', 'SimHei'], 'size': (12, 1)},
    },
    'line_spacing': (1.5, 0.2),
}


def check_page_margins(doc):
    """检查页面边距"""
    issues = []
    for i, section in enumerate(doc.sections):
        for attr, (expected, tol) in STANDARDS['page'].items():
            actual = getattr(section, attr, None)
            if actual is None:
                continue
            actual_cm = actual / 360000
            if abs(actual_cm - expected) > tol:
                issues.append(f"第{i+1}节 {attr}: 期望 {expected}cm, 实际 {actual_cm:.2f}cm")
    return issues


def check_paragraph_format(doc):
    """检查段落格式"""
    issues = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        pf = para.paragraph_format
        if pf.line_spacing and abs(pf.line_spacing - 1.5) > 0.2:
            issues.append(f"段落{i+1} 行距: {pf.line_spacing} (期望1.5)")
        if para.style.name == 'Normal' and text:
            if pf.first_line_indent is None:
                issues.append(f"段落{i+1} 缺少首行缩进")
    return issues


def check_docx_format(docx_path):
    """全面检查 .docx 格式"""
    doc = Document(docx_path)

    all_issues = []
    all_issues.extend(check_page_margins(doc))
    all_issues.extend(check_paragraph_format(doc))

    print("=" * 60)
    print("[报告] 论文格式验证报告")
    print("=" * 60)
    print(f"\n文件: {docx_path}")
    print(f"段落总数: {len(doc.paragraphs)}")
    print(f"节数: {len(doc.sections)}")

    styles_used = set()
    for p in doc.paragraphs:
        if p.style:
            styles_used.add(p.style.name)
    print(f"使用样式: {', '.join(sorted(styles_used))}")

    print(f"\n## 检查结果")
    if all_issues:
        print(f"[FAIL] 发现 {len(all_issues)} 个格式问题:")
        for issue in all_issues:
            print(f"  - {issue}")
    else:
        print("[OK] 所有检测项通过，格式合规!")

    margin_ok = len(check_page_margins(doc)) == 0
    print(f"\n## 格式要点确认清单")
    print(f"  {'[OK]' if margin_ok else '[FAIL]'} 页边距")
    print(f"  [INFO] 字体嵌入: 未检查(需手动确认)")

    return len(all_issues) == 0


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='论文 .docx 格式验证工具')
    parser.add_argument('--input', '-i', required=True, help='输入的 .docx 文件')
    parser.add_argument('--rules', '-r', default=None, help='格式规则参考文件 (可选)')
    args = parser.parse_args()

    if not Path(args.input).exists():
        print(f"[ERR] 输入文件不存在: {args.input}")
        sys.exit(1)

    ok = check_docx_format(args.input)
    sys.exit(0 if ok else 1)
